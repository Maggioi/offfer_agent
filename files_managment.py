import openpyxl
import msoffcrypto
from datetime import datetime
import io
import os
import docx

def open_calculator(nazwa_oferty):
    '''Opening a passworded calculator'''

    with open("password.txt", "r") as f:
        password = f.read()
    decrypted_workbook = io.BytesIO()

    with open(nazwa_oferty, "rb") as encrypted_workbook:
        file = msoffcrypto.OfficeFile(encrypted_workbook)
        file.load_key(password=password)
        file.decrypt(decrypted_workbook)

    wb = openpyxl.load_workbook(decrypted_workbook, data_only=True)
    return wb   
    
def open_offer(doc, tabela, data, nr_sciany):
    '''Filling wall table with informations'''

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            # Dating the offer
            if "123-MM-26" in run.text:
                    run.text = run.text.replace("123-MM-26", data["nr_oferty"])
                    run.text = run.text.replace("02.02.2026r.", f"{datetime.now().strftime("%d.%m.%Y")}r")


    for cell in doc.tables[0]._cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    # Naming the client and the project
                    if "klient" in run.text:
                        run.text = run.text.replace("klient", data["klient"])
                    if "projekt" in run.text:
                        run.text = run.text.replace("projekt", data["projekt"])
                

    for cell in tabela._cells:
        for paragraph in cell.paragraphs:
            # Check the text at the paragraph level first to avoid unnecessary iteration over runs
            p_text = paragraph.text
            
            # If the paragraph is empty, skip to the next one
            if not p_text:
                continue

            # For all other standard replacements, safely iterate through the text runs
            for run in paragraph.runs:
                if not run.text:
                    continue

                if "Wx" in run.text:
                    
                    run.text = run.text.replace("Wx", nr_sciany)
                
                if "certyfikacja" in run.text:
                    run.text = run.text.replace("certyfikacja", data["certyfikacja_BRI"])
                
                if "ile" in run.text:
                    run.text = run.text.replace("ile", data["system"])

                # Safe replacement for Dimensions (L)
                if "Podaj długość" in run.text:
                    run.text = run.text.replace("Podaj długość", data["dlugosc"])
                    
                # Safe replacement for Dimensions (H)
                if "Podaj wysokość" in run.text:
                    run.text = run.text.replace("Podaj wysokość", data["wysokosc"])

                # Safe replacement for Parking / Suspension setup
                if "Kac" in run.text:
                    run.text = run.text.replace("Kac", data["parkowanie"])
                   
                # Replacement for the number of modules
                if "Liczba modułów" in run.text:
                    run.text = run.text.replace("Liczba modułów", data["liczba_modulow"])
                
                if "akustyka" in run.text:
                    run.text = run.text.replace("akustyka", data["akustyka"])

                # Replacement for the board/panel type
                if "pyta" in run.text:
                    run.text = run.text.replace("pyta", data["plyta"])

                # Replacement for the operation type
                if "operacja" in run.text:
                    run.text = run.text.replace("operacja", data["polautomat"])
                
                # Replacement for the track color
                if "tory" in run.text:
                    run.text = run.text.replace("tory", data["kolor_toru"])
                
                # Replacement for the door type
                if "drzwi" in run.text:
                    run.text = run.text.replace("drzwi", data["drzwi"])

                # Replacement for the glass type
                if "szkło" in run.text:
                    run.text = run.text.replace("szkło", data["szklo"])
                
                # Replacement for the profile powder coating / finish
                if "aluminium" in run.text:
                    run.text = run.text.replace("aluminium", data["lakierowane_profile"])
                
                # Replacement for additional equipment:
                if "adyszynal" in run.text:
                    run.text = run.text.replace("adyszynal", data["additional_equipment"])

                # Replacement for the price 
                if "Cena" in run.text:
                    run.text = run.text.replace("Cena", data["cena"])
    return doc

def update_summary_table(doc, calkowita_liczba_scian, cena_wszystkich, additional_items, doplaty):
    """Updating the last table"""
    summary_table = doc.tables[-3]

    if len(additional_items) > 0:
        old_line = summary_table.rows[-1]
        for _ in range(len(additional_items)):
            new_line = summary_table.add_row()
            new_line.cells[1].merge(new_line.cells[2])
            new_line.cells[0].text = f"{additional_items[_]}:"
            new_line.cells[1].text = f"+ {doplaty[_]}"
            new_line.height = old_line.height

            
            for cell in old_line.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        # copying font from the row above to paste it later
                        if "Cena wszystkich" in run.text:
                            font_cena = run.font.name
                            font_size_cena = run.font.size

            # CELL FORMATTING
            i = 0
            for cell in new_line.cells:
                cell.vertical_alignment = docx.enum.table.WD_ALIGN_VERTICAL.CENTER
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                            run.font.name = font_cena
                            run.font.size = font_size_cena
                            if i == 1:
                                run.bold = True
                                paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER                   
                i += 1
    
    for cell in summary_table._cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:

            #if calkowita_liczba_scian > 1:
                if "Cena wszystkich" in run.text:
                    if calkowita_liczba_scian > 1:
                        run.text = run.text.replace("Cena wszystkich", f"Total {calkowita_liczba_scian} walls")
                    else:
                        run.text = run.text.replace("Cena wszystkich", f"Total {calkowita_liczba_scian} wall")
                if "ewro" in run.text:
                    run.text = run.text.replace("ewro", cena_wszystkich)

            # #else:
            #     if "Cena wszystkich" in run.text:
            #         run.text = run.text.replace("Cena wszystkich:", "")
            #     if "ewro" in run.text:
            #         run.text = run.text.replace("ewro", "")
        

def decrypt_excel_file(source, destination):
    """Decrypting at bytes level, preserving formulas"""

    with open("password.txt", "r") as f:
            password = f.read()
            
    with open(source, "rb") as f:
        office_file = msoffcrypto.OfficeFile(f)
        office_file.load_key(password=password)
        
        with open(destination, "wb") as decrypted:
            office_file.decrypt(decrypted)

def save_offer(doc, nazwa_oferty):
    # 4. Saving filled offer document.
    nazwa_oferty = nazwa_oferty.replace(".xlsx", ".docx")
    doc.save(nazwa_oferty)
    print(f"Sukces! Wygenerowano plik: {nazwa_oferty}")
    os.startfile(nazwa_oferty)